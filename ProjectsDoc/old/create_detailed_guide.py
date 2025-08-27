#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Django Proje Adımları Kılavuzu Oluşturucu
Excel dosyasındaki tüm adımları okuyup detaylı DOCX kılavuzu oluşturur.
"""

import pandas as pd
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
import os
from datetime import datetime

def read_excel_phases():
    """Excel dosyasından proje aşamalarını okur"""
    try:
        # Excel dosyasını oku
        file_path = r"c:\Users\ahmet.yildirir\Desktop\Apps\ProjectsDoc\Django_Proje_Yonetimi_Taslagi.xlsx"
        df = pd.read_excel(file_path, sheet_name='🎯 Detaylı Proje Aşamaları')
        
        phases = []
        for index, row in df.iterrows():
            if pd.notna(row.iloc[0]) and str(row.iloc[0]).isdigit():  # Sıra numarası var mı kontrol et
                phase = {
                    'sira': str(row.iloc[0]).zfill(3),
                    'baslik': str(row.iloc[1]),
                    'aciklama': str(row.iloc[2]),
                    'sure': str(row.iloc[3]),
                    'oncelik': str(row.iloc[4]),
                    'ipuclari': str(row.iloc[8]) if pd.notna(row.iloc[8]) else "",
                    'bagimlilik': str(row.iloc[9]) if pd.notna(row.iloc[9]) else "",
                    'notlar': str(row.iloc[10]) if pd.notna(row.iloc[10]) else "",
                    'kategori': str(row.iloc[11]) if pd.notna(row.iloc[11]) else ""
                }
                phases.append(phase)
        
        return phases
    except Exception as e:
        print(f"Excel okuma hatası: {e}")
        return []

def create_detailed_guide():
    """Detaylı Django proje kılavuzu oluştur"""
    
    # Excel'den aşamaları oku
    phases = read_excel_phases()
    
    if not phases:
        print("❌ Excel dosyasından aşamalar okunamadı!")
        return
    
    # Yeni Word dökümanı oluştur
    doc = Document()
    
    # Ana başlık
    title = doc.add_heading('DJANGO PROJESİ GELİŞTİRME KILAVUZU', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_paragraph('Başlangıçtan Canlı Yayına Kadar Adım Adım Rehber')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].font.size = Pt(16)
    subtitle.runs[0].font.italic = True
    
    # Tarih
    date_para = doc.add_paragraph(f'Oluşturulma Tarihi: {datetime.now().strftime("%d.%m.%Y")}')
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_para.runs[0].font.size = Pt(12)
    
    # Sayfa kırma
    doc.add_page_break()
    
    # İçindekiler (manuel)
    doc.add_heading('İÇİNDEKİLER', level=1)
    toc_items = [
        "1. GİRİŞ VE KAPSAM",
        "2. ÖN GEREKSİNİMLER",
        "3. PROJE AŞAMALARI DETAYLARİ",
        "4. YAYGIN HATALAR VE ÇÖZÜMLER",
        "5. İPUÇLARI VE BEST PRACTICES",
        "6. KAYNAKLAR VE REFERANSLAR"
    ]
    
    for item in toc_items:
        p = doc.add_paragraph(item)
        p.style = 'List Number'
    
    doc.add_page_break()
    
    # Giriş bölümü
    doc.add_heading('1. GİRİŞ VE KAPSAM', level=1)
    
    intro_text = """
Bu kılavuz, Django web framework'ü kullanarak profesyonel seviyede web uygulamaları geliştirmek isteyen geliştiriciler için hazırlanmıştır. 
Başlangıç seviyesinden ileri seviyeye kadar tüm aşamaları kapsar ve real-world projelerde karşılaşabileceğiniz sorunlar için çözümler sunar.

Bu kılavuzun özelliklerl:
• 108 detaylı adım ile kapsamlı kapsam
• Her adım için pratik örnekler
• Yaygın hatalar ve çözüm yolları
• Production deployment rehberi
• Security best practices
• Performance optimization teknikleri
• Tek geliştirici odaklı yaklaşım
"""
    
    doc.add_paragraph(intro_text)
    
    # Ön gereksinimler
    doc.add_heading('2. ÖN GEREKSİNİMLER', level=1)
    
    prereq_text = """
Bu kılavuzu takip etmek için aşağıdaki bilgi ve araçlara sahip olmanız gerekmektedir:

• Temel Python bilgisi (değişkenler, fonksiyonlar, sınıflar)
• HTML, CSS temel bilgisi
• Command line/terminal kullanımı
• Git version control temel bilgisi
• PostgreSQL veritabanı temel kavramları

Gerekli Araçlar:
• Python 3.8 veya üzeri
• Code editor (VS Code önerilir)
• Git
• PostgreSQL
• Web browser
• Terminal/Command Prompt
"""
    
    doc.add_paragraph(prereq_text)
    
    doc.add_page_break()
    
    # Ana bölüm - Proje aşamaları
    doc.add_heading('3. PROJE AŞAMALARI DETAYLARİ', level=1)
    
    # Kategori bazında grupla
    categories = {}
    for phase in phases:
        cat = phase['kategori']
        if cat not in categories:
            categories[cat] = []
        categories[cat].append(phase)
    
    # Her kategori için detaylı açıklama
    category_details = {
        '🎯 Planlama': {
            'baslik': 'PROJE PLANLAMA VE ANALİZ AŞAMASI',
            'aciklama': 'Bu aşamada projenin temellerini atıyoruz. Doğru planlama yapılmayan projeler genellikle başarısız olur.'
        },
        '⚙️ Kurulum': {
            'baslik': 'GELİŞTİRME ORTAMI KURULUMU',
            'aciklama': 'Development environment kurulumu ve temel araçların yapılandırılması.'
        },
        '🗄️ Veritabanı': {
            'baslik': 'VERİTABANI TASARIMI VE MODEL OLUŞTURMA',
            'aciklama': 'Django ORM kullanarak veritabanı modellerinin tasarlanması ve oluşturulması.'
        },
        '🌐 Backend': {
            'baslik': 'BACKEND GELİŞTİRME',
            'aciklama': 'Django views, URLs ve API endpoint\'lerinin geliştirilmesi.'
        },
        '🎨 Frontend': {
            'baslik': 'FRONTEND GELİŞTİRME',
            'aciklama': 'Template\'ler, static files ve kullanıcı arayüzünün geliştirilmesi.'
        },
        '📝 Forms': {
            'baslik': 'FORM YÖNETİMİ VE VALİDASYON',
            'aciklama': 'Django forms kullanarak veri girişi ve validation işlemlerinin yapılandırılması.'
        },
        '📁 Static': {
            'baslik': 'STATIC FILES VE MEDIA YÖNETİMİ',
            'aciklama': 'CSS, JavaScript ve media dosyalarının yönetimi.'
        },
        '🔒 Security': {
            'baslik': 'GÜVENLİK VE İZİN YÖNETİMİ',
            'aciklama': 'Uygulama güvenliği ve kullanıcı yetki sistemlerinin kurulması.'
        },
        '🧪 Testing': {
            'baslik': 'TEST YAZMA VE KALİTE KONTROL',
            'aciklama': 'Unit test, integration test ve kalite kontrol süreçlerinin uygulanması.'
        },
        '⚡ Performance': {
            'baslik': 'PERFORMANS OPTİMİZASYONU',
            'aciklama': 'Uygulama performansının artırılması ve optimizasyon teknikleri.'
        },
        '🚀 Deployment': {
            'baslik': 'DEPLOYMENT VE CANLI YAYIN',
            'aciklama': 'Uygulamanın production ortamına deploy edilmesi.'
        },
        '📊 Monitoring': {
            'baslik': 'MONİTORING VE LOGGİNG',
            'aciklama': 'Uygulama izleme, log yönetimi ve error tracking.'
        },
        '📚 Documentation': {
            'baslik': 'DOKÜMANTASYON',
            'aciklama': 'Teknik ve kullanıcı dokümantasyonunun hazırlanması.'
        },
        '✅ QA': {
            'baslik': 'KALİTE GÜVENCESİ VE TEST',
            'aciklama': 'Kapsamlı test süreçleri ve kalite kontrol işlemleri.'
        },
        '🎯 Launch': {
            'baslik': 'PROJE TAMAMLAMA VE LAUNCH',
            'aciklama': 'Projenin tamamlanması ve canlı yayına çıkış süreçleri.'
        },
        '🔧 Maintenance': {
            'baslik': 'BAKIM VE DESTEK',
            'aciklama': 'Proje sonrası bakım, güncelleme ve destek süreçleri.'
        }
    }
    
    for category, category_phases in categories.items():
        if category in category_details:
            # Kategori başlığı
            doc.add_heading(f"{category_details[category]['baslik']}", level=2)
            doc.add_paragraph(category_details[category]['aciklama'])
            
            # Bu kategorideki her adım
            for phase in category_phases:
                create_phase_detail(doc, phase)
    
    # Ek bölümler
    add_common_errors_section(doc)
    add_tips_section(doc)
    add_resources_section(doc)
    
    # Dosyayı kaydet
    output_path = r"c:\Users\ahmet.yildirir\Desktop\Apps\ProjectsDoc\Django_Detayli_Proje_Kilavuzu.docx"
    doc.save(output_path)
    print(f"✅ Detaylı kılavuz oluşturuldu: {output_path}")

def create_phase_detail(doc, phase):
    """Her bir aşama için detaylı açıklama oluştur"""
    
    # Aşama başlığı
    heading = f"ADIM {phase['sira']}: {phase['baslik']}"
    doc.add_heading(heading, level=3)
    
    # Temel bilgiler tablosu
    table = doc.add_table(rows=5, cols=2)
    table.style = 'Table Grid'
    
    table.cell(0, 0).text = "Öncelik"
    table.cell(0, 1).text = phase['oncelik']
    table.cell(1, 0).text = "Tahmini Süre"
    table.cell(1, 1).text = phase['sure']
    table.cell(2, 0).text = "Bağımlılıklar"
    table.cell(2, 1).text = phase['bagimlilik']
    table.cell(3, 0).text = "Kategori"
    table.cell(3, 1).text = phase['kategori']
    table.cell(4, 0).text = "İpuçları"
    table.cell(4, 1).text = phase['ipuclari']
    
    # Detaylı açıklama
    doc.add_paragraph("AÇIKLAMA:", style='Heading 4')
    doc.add_paragraph(phase['aciklama'])
    
    # Aşamaya göre özel detaylar ekle
    add_specific_details(doc, phase)
    
    # Ayırıcı çizgi
    doc.add_paragraph("─" * 80)

def add_specific_details(doc, phase):
    """Her aşama için özel detaylar ekle"""
    
    title = phase['baslik'].lower()
    
    if "proje fikrini netleştir" in title:
        add_project_idea_details(doc)
    elif "gereksinim analizi" in title:
        add_requirement_analysis_details(doc)
    elif "python kurulumu" in title:
        add_python_installation_details(doc)
    elif "virtual environment" in title:
        add_venv_details(doc)
    elif "django projesi oluştur" in title:
        add_django_project_details(doc)
    elif "model" in title:
        add_model_details(doc)
    elif "migration" in title:
        add_migration_details(doc)
    elif "view" in title:
        add_view_details(doc)
    elif "template" in title:
        add_template_details(doc)
    elif "form" in title:
        add_form_details(doc)
    elif "test" in title:
        add_test_details(doc)
    elif "deployment" in title:
        add_deployment_details(doc)
    else:
        add_generic_details(doc, phase)

def add_project_idea_details(doc):
    """Proje fikri netleştirme detayları"""
    doc.add_paragraph("NASIL YAPILIR:", style='Heading 4')
    
    details = """
1. Problemi Tanımlayın:
   • Hangi sorunu çözüyorsunuz?
   • Bu sorun gerçekten var mı?
   • İnsanlar bunun için para ödeyebilir mi?

2. Hedef Kitle Analizi:
   • Uygulamanızı kimler kullanacak?
   • Yaş, meslek, teknik seviye?
   • Hangi platformları tercih ediyorlar?

3. Rakip Analizi:
   • Benzer uygulamalar var mı?
   • Onlardan farkınız ne?
   • Nasıl daha iyi olabilirsiniz?

4. Özellik Listesi:
   • Temel özellikler (MVP)
   • Gelecekteki özellikler
   • Nice-to-have özellikler

5. Başarı Metrikleri:
   • Hangi sayılar başarıyı gösterecek?
   • Kullanıcı sayısı, işlem hacmi, gelir?
"""
    
    doc.add_paragraph(details)
    
    doc.add_paragraph("ÖRNEK PROJE FİKRİ:", style='Heading 4')
    example = """
Proje: "Öğrenci Yemekhane Sistemi"
Problem: Öğrenciler yemekhane menüsünü göremiyorlar ve rezervasyon yapamiyorlar.
Hedef Kitle: Üniversite öğrencileri (18-25 yaş)
Temel Özellikler: Menü görüntüleme, rezervasyon, ödeme, admin paneli
Başarı Metriği: Günlük 500+ aktif kullanıcı
"""
    doc.add_paragraph(example)

def add_requirement_analysis_details(doc):
    """Gereksinim analizi detayları"""
    doc.add_paragraph("NASIL YAPILIR:", style='Heading 4')
    
    details = """
1. Fonksiyonel Gereksinimler:
   • Kullanıcı neler yapabilecek?
   • Sistemin hangi işlevleri olacak?
   • Input ve output'lar neler?

2. Teknik Gereksinimler:
   • Hangi teknolojiler kullanılacak?
   • Performans beklentileri?
   • Güvenlik gereksinimleri?

3. User Story Yazma:
   Format: "Ben [kullanıcı tipi] olarak, [hedef] için [özellik] istiyorum."

4. Acceptance Criteria:
   • Her user story için kabul kriterleri
   • Test edilebilir koşullar

5. Non-Functional Requirements:
   • Performance (yanıt süresi)
   • Security (veri korunması)
   • Usability (kullanım kolaylığı)
   • Scalability (ölçeklenebilirlik)
"""
    
    doc.add_paragraph(details)
    
    doc.add_paragraph("ÖRNEK USER STORY:", style='Heading 4')
    example = """
"Ben öğrenci olarak, bugünkü yemek menüsünü görmek için uygulamaya giriş yapmak istiyorum."

Acceptance Criteria:
✓ Kullanıcı email/şifre ile giriş yapabilmeli
✓ Giriş sonrası anasayfada günün menüsü görünmeli
✓ Menüde yemek adı, fiyat ve kalori bilgisi olmalı
✓ Yanlış şifre girişinde hata mesajı gösterilmeli
"""
    doc.add_paragraph(example)

def add_python_installation_details(doc):
    """Python kurulum detayları"""
    doc.add_paragraph("NASIL YAPILIR:", style='Heading 4')
    
    details = """
1. Python Sürümü Kontrolü:
   Terminalde: python --version veya python3 --version
   
2. Windows için Kurulum:
   • python.org'dan Python 3.8+ indirin
   • "Add to PATH" seçeneğini işaretleyin
   • pip otomatik olarak gelir

3. macOS için Kurulum:
   • Homebrew ile: brew install python3
   • Veya python.org'dan installer

4. Linux için Kurulum:
   • Ubuntu/Debian: sudo apt-get install python3 python3-pip
   • CentOS/RHEL: sudo yum install python3 python3-pip

5. Kurulum Testi:
   python --version
   pip --version
   
6. Pip Güncellemesi:
   python -m pip install --upgrade pip
"""
    
    doc.add_paragraph(details)
    
    doc.add_paragraph("YAYGIN SORUNLAR:", style='Heading 4')
    problems = """
❌ Problem: "python command not found"
✅ Çözüm: PATH değişkenine Python ekleyin

❌ Problem: Python 2.x kurulu ama 3.x gerekli  
✅ Çözüm: python3 komutunu kullanın veya alias oluşturun

❌ Problem: pip çalışmıyor
✅ Çözüm: python -m pip komutunu kullanın
"""
    doc.add_paragraph(problems)

def add_venv_details(doc):
    """Virtual environment detayları"""
    doc.add_paragraph("NASIL YAPILIR:", style='Heading 4')
    
    details = """
1. Virtual Environment Oluşturma:
   Windows: python -m venv venv
   macOS/Linux: python3 -m venv venv

2. Aktivasyon:
   Windows: venv\\Scripts\\activate
   macOS/Linux: source venv/bin/activate

3. Aktivasyon Kontrolü:
   Terminal'de (venv) yazısını görmelisiniz
   which python komutu venv içindeki python'u göstermeli

4. Deaktivasyon:
   deactivate komutu

5. Paket Yükleme:
   pip install django
   pip install psycopg2-binary
   
6. Requirements Dosyası:
   pip freeze > requirements.txt
   pip install -r requirements.txt
"""
    
    doc.add_paragraph(details)
    
    doc.add_paragraph("NEDEN GEREKLİ:", style='Heading 4')
    why = """
• Global Python'u kirletmez
• Proje bağımlılıklarını izole eder
• Farklı projeler farklı paket versiyonları kullanabilir
• Production deployment kolaylaşır
• Takım üyeleri aynı environment'ı kullanabilir
"""
    doc.add_paragraph(why)

def add_django_project_details(doc):
    """Django proje oluşturma detayları"""
    doc.add_paragraph("NASIL YAPILIR:", style='Heading 4')
    
    details = """
1. Django Kurulumu:
   pip install django

2. Proje Oluşturma:
   django-admin startproject myproject
   
3. Klasör Yapısı İncelemesi:
   myproject/
   ├── manage.py (Django komutları için)
   ├── myproject/
   │   ├── __init__.py
   │   ├── settings.py (ayarlar)
   │   ├── urls.py (URL routing)
   │   ├── wsgi.py (deployment için)
   │   └── asgi.py (async için)

4. İlk Çalıştırma:
   cd myproject
   python manage.py runserver
   
5. Tarayıcıda Test:
   http://127.0.0.1:8000/
   Django welcome sayfasını görmelisiniz

6. Admin Superuser:
   python manage.py createsuperuser
"""
    
    doc.add_paragraph(details)
    
    doc.add_paragraph("PROJE ADI SEÇİMİ:", style='Heading 4')
    naming = """
✅ İyi Örnekler: blog, ecommerce, taskmanager
❌ Kötü Örnekler: mysite, project, django_app

Kurallar:
• Küçük harflerle
• Alt çizgi kullanabilirsiniz
• Türkçe karakter yok
• Python keyword'leri kullanmayın (if, for, class vb.)
• Django'nun builtin modül adlarını kullanmayın
"""
    doc.add_paragraph(naming)

def add_model_details(doc):
    """Model oluşturma detayları"""
    doc.add_paragraph("NASIL YAPILIR:", style='Heading 4')
    
    details = """
1. Model Sınıfı Tanımlama:
   
   class User(models.Model):
       username = models.CharField(max_length=150, unique=True)
       email = models.EmailField()
       first_name = models.CharField(max_length=30, blank=True)
       is_active = models.BooleanField(default=True)
       created_at = models.DateTimeField(auto_now_add=True)
       
       class Meta:
           db_table = 'users'
           ordering = ['-created_at']
           
       def __str__(self):
           return self.username

2. Field Tipleri:
   • CharField: Kısa metin (max_length zorunlu)
   • TextField: Uzun metin
   • IntegerField: Tam sayı
   • BooleanField: True/False
   • DateTimeField: Tarih ve saat
   • EmailField: Email formatı
   • ForeignKey: İlişki (başka model)

3. Field Parametreleri:
   • null=True: Database'de NULL olabilir
   • blank=True: Form'da boş geçilebilir
   • default: Varsayılan değer
   • unique=True: Benzersiz olmalı
   • db_index=True: Index oluştur
"""
    
    doc.add_paragraph(details)

def add_migration_details(doc):
    """Migration detayları"""
    doc.add_paragraph("NASIL YAPILIR:", style='Heading 4')
    
    details = """
1. Migration Dosyası Oluşturma:
   python manage.py makemigrations
   
2. Migration Dosyasını İnceleme:
   migrations/0001_initial.py dosyasını açın
   Ne değiştiğini kontrol edin

3. Migration Uygulama:
   python manage.py migrate

4. Migration Durumu:
   python manage.py showmigrations

5. Belirli App için Migration:
   python manage.py makemigrations myapp
   python manage.py migrate myapp

6. Migration Geri Alma:
   python manage.py migrate myapp 0001

7. Fake Migration:
   python manage.py migrate --fake
"""
    
    doc.add_paragraph(details)
    
    doc.add_paragraph("ÖNEMLİ UYARILAR:", style='Heading 4')
    warnings = """
⚠️ Migration dosyalarını silmeyin!
⚠️ Migration dosyalarını manuel düzenlemeyin!
⚠️ Production'da migration öncesi backup alın!
⚠️ Migration conflict'ları dikkatli çözün!
⚠️ Büyük tablolarda downtime olabileceğini düşünün!
"""
    doc.add_paragraph(warnings)

def add_view_details(doc):
    """View oluşturma detayları"""
    doc.add_paragraph("NASIL YAPILIR:", style='Heading 4')
    
    details = """
1. Function-Based View:
   
   def home(request):
       if request.method == 'POST':
           # POST işlemleri
           pass
       
       context = {
           'title': 'Ana Sayfa',
           'users': User.objects.all()
       }
       return render(request, 'home.html', context)

2. Class-Based View:
   
   class HomeView(TemplateView):
       template_name = 'home.html'
       
       def get_context_data(self, **kwargs):
           context = super().get_context_data(**kwargs)
           context['users'] = User.objects.all()
           return context

3. URL Bağlama (urls.py):
   
   urlpatterns = [
       path('', views.home, name='home'),
       path('home/', HomeView.as_view(), name='home_class'),
   ]

4. Generic Views:
   • ListView: Liste görünümü
   • DetailView: Detay görünümü  
   • CreateView: Oluşturma formu
   • UpdateView: Güncelleme formu
   • DeleteView: Silme onayı
"""
    
    doc.add_paragraph(details)

def add_template_details(doc):
    """Template oluşturma detayları"""
    doc.add_paragraph("NASIL YAPILIR:", style='Heading 4')
    
    details = """
1. Template Klasör Yapısı:
   templates/
   ├── base.html
   ├── home.html
   └── users/
       ├── list.html
       └── detail.html

2. Base Template (base.html):
   
   <!DOCTYPE html>
   <html>
   <head>
       <title>{% block title %}My Site{% endblock %}</title>
   </head>
   <body>
       <nav>
           <!-- Navigation -->
       </nav>
       
       <main>
           {% block content %}
           {% endblock %}
       </main>
   </body>
   </html>

3. Child Template:
   
   {% extends 'base.html' %}
   
   {% block title %}Home Page{% endblock %}
   
   {% block content %}
       <h1>Welcome!</h1>
       {% for user in users %}
           <p>{{ user.username }}</p>
       {% endfor %}
   {% endblock %}

4. Template Tags:
   • {{ variable }}: Değişken yazdırma
   • {% tag %}: Mantık bloğu
   • {% comment %}: Yorum
   • {% load static %}: Static dosyalar
"""
    
    doc.add_paragraph(details)

def add_form_details(doc):
    """Form oluşturma detayları"""
    doc.add_paragraph("NASIL YAPILIR:", style='Heading 4')
    
    details = """
1. Model Form:
   
   class UserForm(forms.ModelForm):
       class Meta:
           model = User
           fields = ['username', 'email', 'first_name']
           widgets = {
               'username': forms.TextInput(attrs={'class': 'form-control'}),
               'email': forms.EmailInput(attrs={'class': 'form-control'}),
           }

2. View'da Form Kullanımı:
   
   def create_user(request):
       if request.method == 'POST':
           form = UserForm(request.POST)
           if form.is_valid():
               form.save()
               return redirect('user_list')
       else:
           form = UserForm()
       
       return render(request, 'create_user.html', {'form': form})

3. Template'te Form:
   
   <form method="post">
       {% csrf_token %}
       {{ form.as_p }}
       <button type="submit">Kaydet</button>
   </form>

4. Custom Validation:
   
   def clean_username(self):
       username = self.cleaned_data['username']
       if User.objects.filter(username=username).exists():
           raise forms.ValidationError("Bu kullanıcı adı zaten alınmış!")
       return username
"""
    
    doc.add_paragraph(details)

def add_test_details(doc):
    """Test yazma detayları"""
    doc.add_paragraph("NASIL YAPILIR:", style='Heading 4')
    
    details = """
1. Test Dosyası (tests.py):
   
   from django.test import TestCase
   from django.urls import reverse
   from .models import User

   class UserModelTest(TestCase):
       def setUp(self):
           self.user = User.objects.create(
               username='testuser',
               email='test@example.com'
           )

       def test_user_creation(self):
           self.assertEqual(self.user.username, 'testuser')
           self.assertTrue(self.user.is_active)

       def test_str_method(self):
           self.assertEqual(str(self.user), 'testuser')

2. View Test:
   
   class UserViewTest(TestCase):
       def test_home_page(self):
           response = self.client.get(reverse('home'))
           self.assertEqual(response.status_code, 200)
           self.assertContains(response, 'Welcome')

3. Test Çalıştırma:
   python manage.py test
   python manage.py test myapp
   python manage.py test myapp.tests.UserModelTest

4. Coverage (Test Kapsamı):
   pip install coverage
   coverage run --source='.' manage.py test
   coverage report
   coverage html
"""
    
    doc.add_paragraph(details)

def add_deployment_details(doc):
    """Deployment detayları"""
    doc.add_paragraph("NASIL YAPILIR:", style='Heading 4')
    
    details = """
1. Production Settings:
   
   # settings/production.py
   DEBUG = False
   ALLOWED_HOSTS = ['yourdomain.com', 'www.yourdomain.com']
   
   DATABASES = {
       'default': {
           'ENGINE': 'django.db.backends.postgresql',
           'NAME': os.environ.get('DB_NAME'),
           'USER': os.environ.get('DB_USER'),
           'PASSWORD': os.environ.get('DB_PASSWORD'),
           'HOST': os.environ.get('DB_HOST'),
           'PORT': os.environ.get('DB_PORT'),
       }
   }

2. Static Files:
   python manage.py collectstatic

3. Gunicorn (WSGI Server):
   pip install gunicorn
   gunicorn myproject.wsgi:application

4. Nginx Configuration:
   
   server {
       listen 80;
       server_name yourdomain.com;
       
       location /static/ {
           alias /path/to/static/files/;
       }
       
       location / {
           proxy_pass http://127.0.0.1:8000;
           proxy_set_header Host $host;
           proxy_set_header X-Real-IP $remote_addr;
       }
   }

5. Environment Variables:
   # .env file
   SECRET_KEY=your-secret-key
   DEBUG=False
   DB_NAME=myproject
   DB_USER=postgres
   DB_PASSWORD=password
"""
    
    doc.add_paragraph(details)

def add_generic_details(doc, phase):
    """Genel detaylar"""
    doc.add_paragraph("DİKKAT EDİLECEK NOKTALAR:", style='Heading 4')
    
    generic_tips = f"""
• {phase['notlar']}
• Adımı tamamlamadan sonrakine geçmeyin
• Hata alırsanız önce dokümantasyonu kontrol edin
• Stack Overflow ve Django documentation'ı kullanın
• Backup almayı unutmayın
• Git commit'lerini düzenli yapın
"""
    
    doc.add_paragraph(generic_tips)

def add_common_errors_section(doc):
    """Yaygın hatalar bölümü"""
    doc.add_page_break()
    doc.add_heading('4. YAYGYN HATALAR VE ÇÖZÜMLER', level=1)
    
    errors_text = """
Bu bölümde Django geliştirme sürecinde sık karşılaşılan hatalar ve çözüm yolları yer almaktadır.

IMPORT HATALARI:
❌ ModuleNotFoundError: No module named 'django'
✅ Virtual environment aktif olup olmadığını kontrol edin

❌ ImportError: cannot import name 'User' from 'django.contrib.auth.models'
✅ Import statement'ı kontrol edin: from django.contrib.auth.models import User

VERİTABANI HATALARI:
❌ django.db.utils.OperationalError: no such table
✅ Migration'ları çalıştırın: python manage.py migrate

❌ UNIQUE constraint failed
✅ Duplicate veri var, önceki veriyi silin veya unique=False yapın

URL HATALARI:
❌ NoReverseMatch at / Reverse for 'home' not found
✅ urls.py'da 'home' adında bir URL pattern olup olmadığını kontrol edin

❌ Page not found (404)
✅ URL pattern'i doğru yazılmış mı kontrol edin

TEMPLATE HATALARI:
❌ TemplateDoesNotExist at /
✅ Template dosyasının doğru yerde olup olmadığını kontrol edin

❌ Invalid block tag: 'endfor'
✅ {% for %} ve {% endfor %} eşleşmesini kontrol edin

STATIC FILES HATALARI:
❌ Static files yüklenmiyor
✅ python manage.py collectstatic çalıştırın
✅ STATIC_URL ve STATIC_ROOT ayarlarını kontrol edin

GÜVENLİK HATALARI:
❌ Forbidden (CSRF token missing)
✅ Form'larda {% csrf_token %} eklemeyi unutmayın

❌ DisallowedHost at /
✅ ALLOWED_HOSTS ayarını kontrol edin
"""
    
    doc.add_paragraph(errors_text)

def add_tips_section(doc):
    """İpuçları bölümü"""
    doc.add_page_break()
    doc.add_heading('5. İPUÇLARI VE BEST PRACTICES', level=1)
    
    tips_text = """
Bu bölümde Django development sürecini hızlandıracak ve kaliteyi artıracak ipuçları bulacaksınız.

GELİŞTİRME İPUÇLARI:
• Django Debug Toolbar kullanın - SQL query'leri görebilirsiniz
• Shell kullanmayı öğrenin: python manage.py shell
• Django admin'i customize edin - development sürecini hızlandırır
• Fixture'lar ile test data oluşturun
• Management command'ları yazın tekrarlayan işler için
• django-extensions paketini kurun (shell_plus, reset_db vs.)

KOD KALİTESİ:
• PEP 8 standartlarına uyun
• Black code formatter kullanın
• flake8 ile linting yapın
• Type hints kullanın (Python 3.5+)
• Docstring yazın fonksiyonlar için
• Magic number kullanmayın, constant tanımlayın

PERFORMANs:
• select_related() ve prefetch_related() kullanın N+1 problem için
• Database index'leri ekleyin sık sorgulanan alanlara
• Caching kullanın (Redis önerilir)
• pagination kullanın uzun listeler için
• Static files için CDN kullanın
• Database query'leri optimize edin

GÜVENLİK:
• Hiçbir zaman DEBUG=True ile production'a çıkmayın
• SECRET_KEY'i güvenli tutun ve rotate edin
• HTTPS kullanın production'da
• User input'larını validate edin
• SQL injection'a karşı ORM kullanın
• XSS'e karşı template escaping aktif olsun

DEPLOYMENT:
• Environment variables kullanın hassas bilgiler için
• Separate settings dosyaları kullanın (dev/staging/prod)
• Automated deployment setup edin (CI/CD)
• Database backup stratejinizi belirleyin
• Monitoring ve logging sistemini kurun
• Health check endpoint'i ekleyin

PROJE YÖNETİMİ:
• Git flow kullanın
• Meaningful commit message'ları yazın
• Branch protection rules ekleyin
• Code review sürecini uygulayın
• Documentation güncel tutun
• Issue tracking kullanın (GitHub Issues, Jira)

TEST STRATEGY:
• Test-driven development (TDD) approach'u deneyin
• Unit test coverage'ı %80+ tutun
• Integration test yazın kritik flow'lar için
• Factory boy kullanın test data için
• Continuous integration setup edin
• Load testing yapın production öncesi

ÖĞRENİM KAYNAKLARI:
• Django documentation (docs.djangoproject.com)
• Django Girls Tutorial
• Two Scoops of Django (kitap)
• Classy Class-Based Views (ccbv.co.uk)
• Django packages (djangopackages.org)
• Real Python Django tutorials

VERİMLİLİK:
• IDE shortcuts öğrenin (VS Code/PyCharm)
• Terminal shortcuts kullanın
• Snippet'lar oluşturun sık kullanılan kod parçaları için
• Multiple monitor setup yapın mümkünse
• Pomodoro tekniği kullanın focus için
• Regular break'ler alın burnout'u önlemek için
"""
    
    doc.add_paragraph(tips_text)

def add_resources_section(doc):
    """Kaynaklar bölümü"""
    doc.add_page_break()
    doc.add_heading('6. KAYNAKLAR VE REFERANSLAR', level=1)
    
    resources_text = """
Django geliştirme sürecinde ihtiyaç duyabileceğiniz kaynaklar:

RESMİ DOKÜMANTASYON:
• Django Documentation: https://docs.djangoproject.com/
• Django REST Framework: https://www.django-rest-framework.org/
• Django Tutorial: https://docs.djangoproject.com/en/stable/intro/tutorial01/

ÖĞRETİM KAYNAKLARI:
• Django Girls Tutorial: https://tutorial.djangogirls.org/
• Mozilla Django Tutorial: https://developer.mozilla.org/en-US/docs/Learn/Server-side/Django
• Real Python Django: https://realpython.com/tutorials/django/
• Django for Beginners (kitap): William S. Vincent
• Two Scoops of Django (kitap): Daniel ve Audrey Feldroy

ARAÇLAR VE PAKETLER:
• Django Debug Toolbar: Debug ve profiling
• Django Extensions: Faydalı management command'ları
• Pillow: Image processing
• Celery: Background tasks
• Redis: Caching ve message broker
• Gunicorn: WSGI server
• Nginx: Web server ve reverse proxy

TEST ARAÇLARI:
• Factory Boy: Test data generation
• Pytest-django: Gelişmiş test framework
• Coverage.py: Test coverage ölçümü
• Selenium: Browser automation

DEPLOYMENT PLATFORMLARI:
• Heroku: Kolay deployment
• DigitalOcean: VPS hosting
• AWS: Cloud services
• PythonAnywhere: Python hosting
• Railway: Modern deployment platform

CODE QUALITY:
• Black: Code formatter
• flake8: Linting
• isort: Import sorting
• pre-commit: Git hooks
• mypy: Type checking

VERİTABANI ARAÇLARI:
• PostgreSQL: Production database
• pgAdmin: PostgreSQL GUI
• DBeaver: Universal database tool
• Redis CLI: Redis command line

MONİTORING:
• Sentry: Error tracking
• New Relic: Performance monitoring
• DataDog: Infrastructure monitoring
• Uptime Robot: Uptime monitoring

TOPLULUK KAYNAKLARI:
• Django Forum: https://forum.djangoproject.com/
• r/django (Reddit): https://reddit.com/r/django
• Django Discord: Django community chat
• Stack Overflow: Programming Q&A
• GitHub: Open source Django projects

BLOGLAR VE NEWSLETTER:
• Django News: Weekly newsletter
• Real Python: Python tutorials
• Planet Django: Django blog aggregator
• Django Tricks: Tips and tricks
• Simple is Better Than Complex: Django tutorials

GÜVENLİK KAYNAKLARI:
• OWASP: Web application security
• Django Security Documentation
• Python Security: Security best practices
• Security Headers: HTTP security headers

Bu kaynakları bookmark'layın ve ihtiyaç duyduğunuzda başvurun. Django topluluğu çok yardımseverdir, 
sorularınızı forum ve Discord kanallarında sorabilirsiniz.
"""
    
    doc.add_paragraph(resources_text)

if __name__ == "__main__":
    create_detailed_guide()
