from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_django_instructions_docx():
    """Django Proje Otomatik Oluşturma Talimatnamesi DOCX"""
    
    doc = Document()
    
    # Ana başlık
    title = doc.add_heading('🚀 DJANGO PROJESİ OTOMATIK OLUŞTURMA TALİMATNAMESİ', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Giriş
    intro = doc.add_paragraph('Bu talimatname ile GitHub Copilot\'a Django projenizi baştan sona oluşturması için gerekli tüm bilgileri sağlayabilirsiniz. Aşağıdaki alanları doldurun ve bu dosyayı Copilot\'a vererek projenizin otomatik olarak oluşturulmasını isteyin.')
    
    doc.add_page_break()
    
    # PROJE TEMEL BİLGİLERİ
    doc.add_heading('📋 PROJE TEMEL BİLGİLERİ', level=1)
    
    doc.add_heading('🎯 Proje Tanımı', level=2)
    
    # Tablo oluştur
    table1 = doc.add_table(rows=5, cols=2)
    table1.style = 'Table Grid'
    
    # Tablo içeriği
    cells1 = [
        ['Proje Adı:', '[Projenizin adını yazın - örn: eticaret_sitesi]'],
        ['Proje Açıklaması:', '[Projenin ne yaptığını kısaca açıklayın - örn: Online kitap satış sitesi]'],
        ['Hedef Kitle:', '[Kimler kullanacak - örn: Kitap severler, öğrenciler]'],
        ['Ana Özellikler:', '[Temel özellikleri listeleyin - örn: Ürün katalogu, sepet, ödeme, kullanıcı kayıt]'],
        ['Proje Türü:', '[E-ticaret, Blog, CRM, Portfolio, vb.]']
    ]
    
    for i, (label, description) in enumerate(cells1):
        table1.rows[i].cells[0].text = label
        table1.rows[i].cells[1].text = description
    
    doc.add_heading('🏗️ Teknik Gereksinimler', level=2)
    
    table2 = doc.add_table(rows=6, cols=2)
    table2.style = 'Table Grid'
    
    cells2 = [
        ['Django Versiyonu:', '[Örn: 5.0.x, 4.2.x LTS]'],
        ['Python Versiyonu:', '[Örn: 3.11, 3.12]'],
        ['Veritabanı:', '[PostgreSQL / MySQL / SQLite]'],
        ['Cache:', '[Redis / Memcached / Yok]'],
        ['Frontend Framework:', '[Bootstrap / Tailwind / Custom CSS / React]'],
        ['API:', '[Django REST Framework gerekli mi? Evet/Hayır]']
    ]
    
    for i, (label, description) in enumerate(cells2):
        table2.rows[i].cells[0].text = label
        table2.rows[i].cells[1].text = description
    
    # KULLANICI ARAYÜZÜ VE TASARIM
    doc.add_heading('🎨 KULLANICI ARAYÜZÜ VE TASARIM', level=1)
    
    doc.add_heading('📱 Sayfalar ve İşlevsellik', level=2)
    
    table3 = doc.add_table(rows=4, cols=2)
    table3.style = 'Table Grid'
    
    cells3 = [
        ['Ana Sayfalar:', '[Hangi sayfalar olacak - örn: Ana sayfa, ürün listesi, ürün detay, sepet, ödeme, profil]'],
        ['Kullanıcı Türleri:', '☐ Misafir kullanıcı\\n☐ Kayıtlı kullanıcı\\n☐ Admin/Yönetici\\n☐ Satıcı'],
        ['Özel İşlevler:', '[Ekstra özellikler - örn: Wishlist, yorumlar, puanlama, kupon sistemi, kargo takibi]'],
        ['Navigation Yapısı:', '[Menü yapısı, breadcrumb, arama özellikleri]']
    ]
    
    for i, (label, description) in enumerate(cells3):
        table3.rows[i].cells[0].text = label
        table3.rows[i].cells[1].text = description
    
    doc.add_heading('🎨 Tasarım Tercihleri', level=2)
    
    table4 = doc.add_table(rows=4, cols=2)
    table4.style = 'Table Grid'
    
    cells4 = [
        ['Renk Teması:', '[Örn: Mavi-beyaz, koyu tema, minimalist]'],
        ['Stil:', '[Modern, klasik, minimalist, renkli]'],
        ['Logo/Marka:', '[Var/Yok - varsa açıklama]'],
        ['Layout:', '[Fixed width, fluid, responsive, mobile-first]']
    ]
    
    for i, (label, description) in enumerate(cells4):
        table4.rows[i].cells[0].text = label
        table4.rows[i].cells[1].text = description
    
    # VERİTABANI YAPISI
    doc.add_heading('🗄️ VERİTABANI YAPISI VE MODELLER', level=1)
    
    doc.add_heading('📊 Ana Veri Modelleri', level=2)
    
    table5 = doc.add_table(rows=3, cols=2)
    table5.style = 'Table Grid'
    
    cells5 = [
        ['Gerekli Modeller:', '[Hangi ana modeller olacak - örn: User, Product, Category, Order, Cart]'],
        ['Model İlişkileri:', '[ForeignKey, ManyToMany, OneToOne ilişkileri açıklayın]'],
        ['Özel Alanlar:', '[Özel field\'lar, validation kuralları, meta bilgiler]']
    ]
    
    for i, (label, description) in enumerate(cells5):
        table5.rows[i].cells[0].text = label
        table5.rows[i].cells[1].text = description
    
    doc.add_paragraph('Örnek Model Yapısı:')
    code_para = doc.add_paragraph()
    code_run = code_para.add_run('''
Product Model:
- name: CharField(max_length=200)
- description: TextField
- price: DecimalField(max_digits=10, decimal_places=2)
- stock_quantity: IntegerField
- category: ForeignKey(Category)
- images: ImageField (çoklu resim gerekli mi?)
- tags: ManyToManyField (etiket sistemi)
- created_at: DateTimeField(auto_now_add=True)
- updated_at: DateTimeField(auto_now=True)
- is_active: BooleanField(default=True)
''')
    code_run.font.name = 'Courier New'
    code_run.font.size = Pt(9)
    
    # GÜVENLİK
    doc.add_heading('🔐 GÜVENLİK VE KULLANICI YÖNETİMİ', level=1)
    
    doc.add_heading('👤 Kimlik Doğrulama', level=2)
    
    table6 = doc.add_table(rows=5, cols=2)
    table6.style = 'Table Grid'
    
    cells6 = [
        ['Kayıt Sistemi:', '[Email/Kullanıcı adı ile kayıt]'],
        ['Giriş Seçenekleri:', '[Email/Username + şifre, sosyal medya girişi]'],
        ['Sosyal Giriş:', '[Google, Facebook, GitHub - hangilerini istiyorsanız]'],
        ['Email Doğrulama:', '[Gerekli/Gerekli değil]'],
        ['Şifre Sıfırlama:', '[Email ile/SMS ile]']
    ]
    
    for i, (label, description) in enumerate(cells6):
        table6.rows[i].cells[0].text = label
        table6.rows[i].cells[1].text = description
    
    doc.add_heading('🛡️ Güvenlik Özellikleri', level=2)
    
    security_list = [
        '☐ reCAPTCHA entegrasyonu',
        '☐ Rate Limiting (DDoS koruması)',
        '☐ CSRF Protection',
        '☐ XSS Protection',
        '☐ SQL Injection koruması',
        '☐ Admin panel güvenliği',
        '☐ HTTPS zorunluluğu',
        '☐ Password strength validation',
        '☐ Two-factor authentication',
        '☐ Session security'
    ]
    
    for item in security_list:
        doc.add_paragraph(item, style='List Bullet')
    
    # ÖDEME SİSTEMİ
    doc.add_heading('💳 ÖDEME VE E-TİCARET', level=1)
    
    table7 = doc.add_table(rows=5, cols=2)
    table7.style = 'Table Grid'
    
    cells7 = [
        ['Ödeme Yöntemleri:', '[Kredi kartı, PayPal, Stripe, iyzico, vb.]'],
        ['Para Birimi:', '[TL, USD, EUR]'],
        ['Kargo Sistemi:', '[Kargo entegrasyonu gerekli mi?]'],
        ['Stok Takibi:', '[Otomatik/Manuel stok yönetimi]'],
        ['İndirim Sistemi:', '[Kupon, promosyon kodu sistemi]']
    ]
    
    for i, (label, description) in enumerate(cells7):
        table7.rows[i].cells[0].text = label
        table7.rows[i].cells[1].text = description
    
    # API VE ENTEGRASYON
    doc.add_heading('🔧 ENTEGRASYON VE API', level=1)
    
    table8 = doc.add_table(rows=6, cols=2)
    table8.style = 'Table Grid'
    
    cells8 = [
        ['REST API:', '[Django REST Framework gerekli mi?]'],
        ['API Dokümantasyonu:', '[Swagger/OpenAPI otomatik oluşturulsun mu?]'],
        ['Dış Servisler:', '[Google Maps, AWS S3, SendGrid, vb.]'],
        ['Social Media:', '[Facebook, Twitter, Instagram entegrasyonu]'],
        ['Analytics:', '[Google Analytics, özel tracking]'],
        ['Email Servisi:', '[Gmail SMTP, SendGrid, AWS SES]']
    ]
    
    for i, (label, description) in enumerate(cells8):
        table8.rows[i].cells[0].text = label
        table8.rows[i].cells[1].text = description
    
    # DEPLOYMENT
    doc.add_heading('🚀 DEPLOYMENT VE HOSTING', level=1)
    
    table9 = doc.add_table(rows=6, cols=2)
    table9.style = 'Table Grid'
    
    cells9 = [
        ['Hosting Platform:', '[Heroku, DigitalOcean, AWS, VPS]'],
        ['Domain:', '[Özel domain var mı?]'],
        ['SSL Sertifikası:', '[Let\'s Encrypt, paid certificate]'],
        ['CDN:', '[CloudFlare, AWS CloudFront]'],
        ['CI/CD:', '[GitHub Actions, GitLab CI]'],
        ['Environment:', '[Development, Staging, Production]']
    ]
    
    for i, (label, description) in enumerate(cells9):
        table9.rows[i].cells[0].text = label
        table9.rows[i].cells[1].text = description
    
    # ÖZEL GEREKSİNİMLER
    doc.add_heading('🎯 ÖZEL GEREKSİNİMLER', level=1)
    
    doc.add_paragraph('Projenize özel, yukarıda yer almayan gereksinimleri buraya yazın:')
    
    special_features = [
        '☐ QR kod ile ürün tarama',
        '☐ AI destekli ürün önerisi',
        '☐ Canlı chat sistemi',
        '☐ Video streaming',
        '☐ Dosya upload/download sistemi',
        '☐ Multi-vendor marketplace',
        '☐ Subscription sistemi',
        '☐ Booking/Rezervasyon sistemi',
        '☐ Çok dilli destek',
        '☐ PWA (Progressive Web App)',
        '☐ Push notifications',
        '☐ Offline çalışma'
    ]
    
    for feature in special_features:
        doc.add_paragraph(feature, style='List Bullet')
    
    doc.add_paragraph()
    doc.add_paragraph('Diğer özel gereksinimler:')
    doc.add_paragraph('_' * 80)
    doc.add_paragraph('_' * 80)
    doc.add_paragraph('_' * 80)
    
    # TESLİM FORMATI
    doc.add_heading('📋 PROJE TESLİM FORMATI', level=1)
    
    deliverables = [
        '☐ Tüm kaynak kodlar (GitHub repository)',
        '☐ Requirements.txt dosyası',
        '☐ README.md dosyası (kurulum kılavuzu)',
        '☐ Environment variables şablonu (.env.example)',
        '☐ Database schema ve initial data',
        '☐ Admin kullanıcısı (username/password)',
        '☐ Test verileri ve test kullanıcıları',
        '☐ Deployment kılavuzu',
        '☐ API dokümantasyonu (varsa)',
        '☐ Kullanıcı kılavuzu'
    ]
    
    for item in deliverables:
        doc.add_paragraph(item, style='List Bullet')
    
    # COPILOT TALİMATI
    doc.add_page_break()
    doc.add_heading('🤖 GITHUB COPILOT TALİMATI', level=1)
    
    doc.add_paragraph('Yukarıdaki formu doldurduktan sonra, aşağıdaki talimatı GitHub Copilot\'a verin:')
    
    instruction_para = doc.add_paragraph()
    instruction_run = instruction_para.add_run('''
"Bu Django proje talimatnamesindeki tüm gereksinimlere göre eksiksiz bir Django projesi oluştur.

PROJE YAPISI:
1. Django projesini başlat ve gerekli ayarları yap
2. Belirtilen modelları oluştur ve migration'ları hazırla
3. Admin panel konfigürasyonunu yap
4. Views, URLs ve templates'leri oluştur
5. Forms ve validation'ları ekle
6. Authentication ve authorization sistemini kur
7. Gerekli static files ve CSS'leri hazırla
8. Test dosyalarını oluştur
9. Requirements.txt ve deployment dosyalarını hazırla
10. README.md ve dokümantasyonu oluştur

GEREKSINIMLER:
- Her adımı detaylı olarak açıkla
- Working code'lar üret
- Production-ready olsun
- Secure olsun
- Scalable olsun
- Best practices uygula
- Comprehensive error handling
- Proper logging

DOSYA YAPISI:
- Clean architecture
- Proper separation of concerns
- Modular design
- Easy to maintain
- Well documented"
''')
    instruction_run.font.name = 'Courier New'
    instruction_run.font.size = Pt(10)
    
    # KULLANIM TALİMATLARI
    doc.add_heading('📝 KULLANIM TALİMATLARI', level=1)
    
    doc.add_heading('🎯 Nasıl Kullanılır', level=2)
    
    usage_steps = [
        'Formu Doldurun: Yukarıdaki tüm alanları projenize göre doldurun',
        'Copilot\'a Verin: Doldurduğunuz formu GitHub Copilot\'a yapıştırın',
        'Talimat Ekleyin: "COPILOT TALİMATI" bölümünü de ekleyin',
        'İsteyin: "Bu talimatlara göre Django projemi oluştur" deyin',
        'Test Edin: Oluşturulan projeyi test edin ve eksikleri belirtin'
    ]
    
    for i, step in enumerate(usage_steps, 1):
        doc.add_paragraph(f'{i}. {step}', style='List Number')
    
    doc.add_heading('💡 İpuçları', level=2)
    
    tips = [
        'Belirsiz alanları boş bırakabilirsiniz, Copilot uygun varsayılanları seçer',
        'Karmaşık projeler için adım adım ilerleyin',
        'Her aşamada test edin ve eksikleri tamamlayın',
        'Dokümantasyonu güncel tutun',
        'Güvenlik ayarlarını production\'da mutlaka kontrol edin',
        'Performance testing yapın',
        'Code review sürecini atlamamayın'
    ]
    
    for tip in tips:
        doc.add_paragraph(f'• {tip}', style='List Bullet')
    
    doc.add_heading('⚠️ Önemli Notlar', level=2)
    
    notes = [
        'Bu talimatname kapsamlı projeler için tasarlanmıştır',
        'Küçük projeler için gerekli bölümleri kullanabilirsiniz',
        'Güvenlik ayarlarını production\'da mutlaka yapın',
        'Backup stratejinizi unutmayın',
        'Performance metrikleri izleyin',
        'Kullanıcı geri bildirimlerini toplayın'
    ]
    
    for note in notes:
        doc.add_paragraph(f'• {note}', style='List Bullet')
    
    # Footer
    doc.add_paragraph()
    footer_para = doc.add_paragraph('Son Güncelleme: 27 Ağustos 2025 | Versiyon: 2.0 | GitHub Copilot Assistant')
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Dosyayı kaydet
    doc.save(r'c:\Users\ahmet.yildirir\Desktop\Apps\ProjectsDoc\Django_Proje_Talimatnamesi.docx')
    print("Django Proje Talimatnamesi DOCX dosyası oluşturuldu!")

if __name__ == "__main__":
    create_django_instructions_docx()
