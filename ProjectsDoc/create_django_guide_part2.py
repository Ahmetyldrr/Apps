from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

def create_django_guide_part2():
    """Django Proje Kılavuzu - Bölüm 2: Gereksinim Analizi ve Teknik Planlama"""
    
    doc = Document()
    
    # Ana başlık
    title = doc.add_heading('🚀 DJANGO PROJESİ GELİŞTİRME KILAVUZU - BÖLÜM 2', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Alt başlık
    subtitle = doc.add_heading('📋 Gereksinim Analizi ve Teknik Planlama', level=1)
    
    doc.add_page_break()
    
    # BÖLÜM 2: GEREKSİNİM ANALİZİ
    heading1 = doc.add_heading('📊 2. GEREKSİNİM ANALİZİ', level=1)
    
    doc.add_heading('🎯 Ne Yapacaksın?', level=2)
    p1 = doc.add_paragraph('Projenin tüm fonksiyonel ve teknik gereksinimlerini detaylıca listele. User story\'leri yaz.')
    
    doc.add_heading('📝 Yapılacaklar:', level=3)
    doc.add_paragraph('• Fonksiyonel gereksinimleri listele (kullanıcı ne yapabilecek?)', style='List Bullet')
    doc.add_paragraph('• Teknik gereksinimleri belirle (performans, güvenlik, ölçeklenebilirlik)', style='List Bullet')
    doc.add_paragraph('• User story\'leri yaz (Kullanıcı olarak... istiyorum... çünkü...)', style='List Bullet')
    doc.add_paragraph('• Acceptance criteria\'ları tanımla', style='List Bullet')
    doc.add_paragraph('• Önceliklendirme yap (Must have, Should have, Could have)', style='List Bullet')
    
    doc.add_heading('💡 Örnek Proje: E-Ticaret Sitesi', level=3)
    
    doc.add_heading('🔹 Fonksiyonel Gereksinimler:', level=4)
    doc.add_paragraph('• Kullanıcı kayıt/giriş sistemi', style='List Bullet')
    doc.add_paragraph('• Ürün katalog yönetimi', style='List Bullet')
    doc.add_paragraph('• Sepet işlemleri', style='List Bullet')
    doc.add_paragraph('• Ödeme sistemi entegrasyonu', style='List Bullet')
    doc.add_paragraph('• Sipariş takibi', style='List Bullet')
    doc.add_paragraph('• Admin panel', style='List Bullet')
    doc.add_paragraph('• Ürün arama ve filtreleme', style='List Bullet')
    
    doc.add_heading('🔹 User Story Örnekleri:', level=4)
    story1 = doc.add_paragraph()
    story1.add_run('Story 1: ').bold = True
    story1.add_run('Müşteri olarak, ürünleri kategorilere göre filtreleyebilmek istiyorum, çünkü aradığım ürünü daha hızlı bulmak istiyorum.')
    
    story2 = doc.add_paragraph()
    story2.add_run('Story 2: ').bold = True
    story2.add_run('İşletme sahibi olarak, ürünlerimi kolayca ekleyip düzenleyebilmek istiyorum, çünkü stok güncellemesi sık yapıyorum.')
    
    story3 = doc.add_paragraph()
    story3.add_run('Story 3: ').bold = True
    story3.add_run('Müşteri olarak, siparişlerimi takip edebilmek istiyorum, çünkü ne zaman geleceğini bilmek istiyorum.')
    
    # TEKNOLOJİ STACK BELİRLEME
    doc.add_heading('🛠️ 3. TEKNOLOJİ STACK\'İ BELİRLE', level=1)
    
    doc.add_heading('🎯 Ne Yapacaksın?', level=2)
    p2 = doc.add_paragraph('Projende kullanacağın tüm teknolojileri, framework\'leri ve araçları seç.')
    
    doc.add_heading('📝 Karar Verilecek Teknolojiler:', level=3)
    doc.add_paragraph('• Django Version (LTS önerisi: Django 4.2+)', style='List Bullet')
    doc.add_paragraph('• Python Version (Python 3.9+ önerisi)', style='List Bullet')
    doc.add_paragraph('• Veritabanı (PostgreSQL, MySQL, SQLite)', style='List Bullet')
    doc.add_paragraph('• Frontend (Django Templates, React, Vue.js)', style='List Bullet')
    doc.add_paragraph('• CSS Framework (Bootstrap, Tailwind, Bulma)', style='List Bullet')
    doc.add_paragraph('• Ödeme Sistemi (Stripe, PayPal, İyzico)', style='List Bullet')
    doc.add_paragraph('• Email Servisi (SendGrid, Mailgun, SMTP)', style='List Bullet')
    doc.add_paragraph('• File Storage (Local, AWS S3, Google Cloud)', style='List Bullet')
    
    doc.add_heading('💡 Örnek Proje Teknoloji Seçimi:', level=3)
    
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Kategori'
    hdr_cells[1].text = 'Seçilen Teknoloji'
    hdr_cells[2].text = 'Sebep'
    
    tech_choices = [
        ['Backend Framework', 'Django 4.2 LTS', 'Stabil, güvenli, geniş dokümantasyon'],
        ['Programming Language', 'Python 3.11', 'Modern özellikler, performans'],
        ['Database', 'PostgreSQL 15', 'Güçlü, ölçeklenebilir, JSON desteği'],
        ['Frontend', 'Django Templates + Bootstrap 5', 'Hızlı geliştirme, responsive'],
        ['Payment', 'Stripe API', 'Güvenli, kolay entegrasyon'],
        ['Email', 'Django built-in SMTP', 'Basit proje için yeterli'],
        ['File Storage', 'Local + AWS S3 (production)', 'Maliyet etkin çözüm'],
        ['Task Queue', 'Celery + Redis', 'Background işlemler için'],
    ]
    
    for choice in tech_choices:
        row_cells = table.add_row().cells
        row_cells[0].text = choice[0]
        row_cells[1].text = choice[1]
        row_cells[2].text = choice[2]
    
    # Dosyayı kaydet
    doc.save(r'c:\Users\ahmet.yildirir\Desktop\Apps\ProjectsDoc\Django_Kilavuz_Bolum2.docx')
    print("Django Kılavuzu Bölüm 2 oluşturuldu!")

if __name__ == "__main__":
    create_django_guide_part2()
