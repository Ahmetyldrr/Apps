from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

def create_django_guide_part1():
    """Django Proje Kılavuzu - Bölüm 1: Temel Bilgiler ve Planlama"""
    
    doc = Document()
    
    # Ana başlık
    title = doc.add_heading('🚀 DJANGO PROJESİ GELİŞTİRME KILAVUZU - BÖLÜM 1', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Alt başlık
    subtitle = doc.add_heading('📋 Temel Bilgiler ve Proje Planlama', level=1)
    
    # Giriş
    intro = doc.add_paragraph()
    intro.add_run('Bu kılavuz, Django projelerini baştan sona geliştirmek için detaylı adımları içerir. ').bold = True
    intro.add_run('Her bölüm örnek proje (E-Ticaret Sitesi) üzerinden açıklanmıştır.')
    
    doc.add_page_break()
    
    # BÖLÜM 1: PROJE PLANLAMA
    heading1 = doc.add_heading('📊 1. PROJE FİKRİNİ NETLEŞTİRME', level=1)
    
    doc.add_heading('🎯 Ne Yapacaksın?', level=2)
    p1 = doc.add_paragraph('Projenin temel amacını, hedef kitlesini ve çözeceği problemi net bir şekilde tanımla.')
    
    doc.add_heading('📝 Yapılacaklar:', level=3)
    doc.add_paragraph('• Proje fikrini 1-2 cümle ile özetle', style='List Bullet')
    doc.add_paragraph('• Hedef kitleyi belirle (yaş, ilgi alanları, teknik seviye)', style='List Bullet')
    doc.add_paragraph('• Hangi problemi çözdüğünü yaz', style='List Bullet')
    doc.add_paragraph('• Benzer projeleri araştır', style='List Bullet')
    doc.add_paragraph('• Projenin değer önerisini tanımla', style='List Bullet')
    
    doc.add_heading('💡 Örnek Proje: E-Ticaret Sitesi', level=3)
    example = doc.add_paragraph()
    example.add_run('Proje Fikri: ').bold = True
    example.add_run('Küçük işletmelerin ürünlerini online satabilecekleri, kullanımı kolay e-ticaret platformu')
    
    example2 = doc.add_paragraph()
    example2.add_run('Hedef Kitle: ').bold = True
    example2.add_run('25-45 yaş arası, teknolojiye orta seviye hakim, küçük işletme sahipleri')
    
    example3 = doc.add_paragraph()
    example3.add_run('Problem: ').bold = True
    example3.add_run('Küçük işletmeler pahalı e-ticaret çözümleri yüzünden online satış yapamıyor')
    
    example4 = doc.add_paragraph()
    example4.add_run('Değer Önerisi: ').bold = True
    example4.add_run('Uygun maliyetli, kolay kurulum, Türkçe destekli e-ticaret çözümü')
    
    doc.add_heading('⚠️ Dikkat Edilecekler:', level=3)
    doc.add_paragraph('• Çok geniş kapsamlı proje planlanmamalı (MVP - Minimum Viable Product düşün)', style='List Bullet')
    doc.add_paragraph('• Hedef kitle çok genel olmamalı (herkes değil, belirli bir grup)', style='List Bullet')
    doc.add_paragraph('• Mevcut çözümlerin eksiklerini tespit et', style='List Bullet')
    
    # Dosyayı kaydet
    doc.save(r'c:\Users\ahmet.yildirir\Desktop\Apps\ProjectsDoc\Django_Kilavuz_Bolum1.docx')
    print("Django Kılavuzu Bölüm 1 oluşturuldu!")

if __name__ == "__main__":
    create_django_guide_part1()
