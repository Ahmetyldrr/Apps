from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

def create_django_guide_part3():
    """Django Proje Kılavuzu - Bölüm 3: Veritabanı Tasarımı ve Model Oluşturma"""
    
    doc = Document()
    
    # Ana başlık
    title = doc.add_heading('🚀 DJANGO PROJESİ GELİŞTİRME KILAVUZU - BÖLÜM 3', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Alt başlık
    subtitle = doc.add_heading('🗄️ Veritabanı Tasarımı ve Model Oluşturma', level=1)
    
    doc.add_page_break()
    
    # VERİTABANI ŞEMASI TASARLAMA
    heading1 = doc.add_heading('🗄️ 4. VERİTABANI ŞEMASI TASARLA', level=1)
    
    doc.add_heading('🎯 Ne Yapacaksın?', level=2)
    p1 = doc.add_paragraph('Projenin tüm veri ihtiyaçlarını analiz et, tabloları ve ilişkileri tasarla.')
    
    doc.add_heading('📝 Yapılacaklar:', level=3)
    doc.add_paragraph('• Entity-Relationship Diagram (ERD) çiz', style='List Bullet')
    doc.add_paragraph('• Ana varlıkları (entities) belirle', style='List Bullet')
    doc.add_paragraph('• Tablolar arası ilişkileri tanımla (1:1, 1:N, N:N)', style='List Bullet')
    doc.add_paragraph('• Her tablonun alanlarını listele', style='List Bullet')
    doc.add_paragraph('• Primary key ve foreign key\'leri belirle', style='List Bullet')
    doc.add_paragraph('• Normalizasyon kurallarını uygula', style='List Bullet')
    
    doc.add_heading('💡 Örnek Proje: E-Ticaret Veritabanı', level=3)
    
    doc.add_heading('🔹 Ana Tablolar:', level=4)
    
    # Tablolar tablosu
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Tablo Adı'
    hdr_cells[1].text = 'Açıklama'
    hdr_cells[2].text = 'Ana Alanlar'
    
    tables_data = [
        ['User', 'Kullanıcı bilgileri', 'id, username, email, password, first_name, last_name, is_active'],
        ['Category', 'Ürün kategorileri', 'id, name, slug, description, parent_category'],
        ['Product', 'Ürün bilgileri', 'id, name, slug, description, price, stock, category, image'],
        ['Cart', 'Alışveriş sepeti', 'id, user, created_at, updated_at'],
        ['CartItem', 'Sepet ürünleri', 'id, cart, product, quantity, price'],
        ['Order', 'Siparişler', 'id, user, total_amount, status, created_at, shipping_address'],
        ['OrderItem', 'Sipariş kalemleri', 'id, order, product, quantity, price'],
        ['Address', 'Adres bilgileri', 'id, user, title, address_line, city, country, postal_code'],
    ]
    
    for table_info in tables_data:
        row_cells = table.add_row().cells
        row_cells[0].text = table_info[0]
        row_cells[1].text = table_info[1]
        row_cells[2].text = table_info[2]
    
    doc.add_heading('🔹 İlişki Türleri:', level=4)
    doc.add_paragraph('• User → Address (1:N) - Bir kullanıcının birden fazla adresi olabilir', style='List Bullet')
    doc.add_paragraph('• Category → Product (1:N) - Bir kategoride birden fazla ürün olabilir', style='List Bullet')
    doc.add_paragraph('• User → Cart (1:1) - Her kullanıcının bir sepeti var', style='List Bullet')
    doc.add_paragraph('• Cart → CartItem (1:N) - Sepette birden fazla ürün olabilir', style='List Bullet')
    doc.add_paragraph('• User → Order (1:N) - Kullanıcının birden fazla siparişi olabilir', style='List Bullet')
    doc.add_paragraph('• Order → OrderItem (1:N) - Siparişte birden fazla ürün olabilir', style='List Bullet')
    
    # DJANGO MODELLERİ OLUŞTURMA
    doc.add_heading('🏗️ 5. DJANGO MODELLERİNİ OLUŞTUR', level=1)
    
    doc.add_heading('🎯 Ne Yapacaksın?', level=2)
    p2 = doc.add_paragraph('Tasarladığın veritabanı şemasını Django modellerine dönüştür.')
    
    doc.add_heading('📝 Yapılacaklar:', level=3)
    doc.add_paragraph('• Her app için models.py dosyasında model sınıflarını yaz', style='List Bullet')
    doc.add_paragraph('• Doğru field tiplerini seç', style='List Bullet')
    doc.add_paragraph('• İlişkileri doğru şekilde tanımla (ForeignKey, ManyToMany, OneToOne)', style='List Bullet')
    doc.add_paragraph('• Meta sınıflarını ekle (verbose_name, ordering, db_table)', style='List Bullet')
    doc.add_paragraph('• __str__ metodlarını yaz', style='List Bullet')
    doc.add_paragraph('• Validation metodları ekle', style='List Bullet')
    
    doc.add_heading('💡 Örnek Model: Product', level=3)
    
    code_para = doc.add_paragraph()
    code_para.add_run('''
from django.db import models
from django.urls import reverse
from django.contrib.auth.models import User

class Category(models.Model):
    name = models.CharField(max_length=100, verbose_name="Kategori Adı")
    slug = models.SlugField(max_length=100, unique=True)
    description = models.TextField(blank=True, verbose_name="Açıklama")
    parent = models.ForeignKey('self', on_delete=models.CASCADE, 
                              null=True, blank=True, 
                              related_name='children')
    created_at = models.DateTimeField(auto_now_add=True)
    
    class Meta:
        verbose_name = "Kategori"
        verbose_name_plural = "Kategoriler"
        ordering = ['name']
        
    def __str__(self):
        return self.name

class Product(models.Model):
    STOCK_STATUS = [
        ('in_stock', 'Stokta'),
        ('out_of_stock', 'Stok Yok'),
        ('discontinued', 'Üretim Durduruldu'),
    ]
    
    name = models.CharField(max_length=200, verbose_name="Ürün Adı")
    slug = models.SlugField(max_length=200, unique=True)
    description = models.TextField(verbose_name="Ürün Açıklaması")
    price = models.DecimalField(max_digits=10, decimal_places=2, 
                               verbose_name="Fiyat")
    stock_quantity = models.PositiveIntegerField(default=0, 
                                               verbose_name="Stok Miktarı")
    stock_status = models.CharField(max_length=20, 
                                  choices=STOCK_STATUS,
                                  default='in_stock')
    category = models.ForeignKey(Category, on_delete=models.CASCADE,
                               related_name='products')
    image = models.ImageField(upload_to='products/', 
                            blank=True, null=True)
    created_at = models.DateTimeField(auto_now_add=True)
    updated_at = models.DateTimeField(auto_now=True)
    is_active = models.BooleanField(default=True)
    
    class Meta:
        verbose_name = "Ürün"
        verbose_name_plural = "Ürünler"
        ordering = ['-created_at']
        indexes = [
            models.Index(fields=['slug']),
            models.Index(fields=['category']),
            models.Index(fields=['price']),
        ]
    
    def __str__(self):
        return self.name
    
    def get_absolute_url(self):
        return reverse('product_detail', args=[self.slug])
    
    def is_in_stock(self):
        return self.stock_quantity > 0 and self.stock_status == 'in_stock'
''').font.name = 'Courier New'
    
    # Dosyayı kaydet
    doc.save(r'c:\Users\ahmet.yildirir\Desktop\Apps\ProjectsDoc\Django_Kilavuz_Bolum3.docx')
    print("Django Kılavuzu Bölüm 3 oluşturuldu!")

if __name__ == "__main__":
    create_django_guide_part3()
