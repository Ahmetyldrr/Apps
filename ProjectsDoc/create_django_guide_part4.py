from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

def create_django_guide_part4():
    """Django Proje Kılavuzu - Bölüm 4: Uygulama ve Deployment"""
    
    doc = Document()
    
    # Ana başlık
    title = doc.add_heading('🚀 DJANGO PROJESİ GELİŞTİRME KILAVUZU - BÖLÜM 4', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Alt başlık
    subtitle = doc.add_heading('🎯 Uygulama Geliştirme ve Deployment', level=1)
    
    doc.add_page_break()
    
    # VIEWS VE URL YAPISI
    heading1 = doc.add_heading('🔗 6. VIEWS VE URL YAPISINI KURGULA', level=1)
    
    doc.add_heading('🎯 Ne Yapacaksın?', level=2)
    p1 = doc.add_paragraph('Kullanıcı isteklerini karşılayacak view fonksiyonları ve URL yapısını oluştur.')
    
    doc.add_heading('📝 Yapılacaklar:', level=3)
    doc.add_paragraph('• URL patterns tasarla (ana site, admin, api)', style='List Bullet')
    doc.add_paragraph('• Function-based veya Class-based views seç', style='List Bullet')
    doc.add_paragraph('• View fonksiyonlarını yaz', style='List Bullet')
    doc.add_paragraph('• URL dispatcher oluştur', style='List Bullet')
    doc.add_paragraph('• Template yapısını planla', style='List Bullet')
    doc.add_paragraph('• Forms oluştur', style='List Bullet')
    
    doc.add_heading('💡 Örnek URL Yapısı:', level=3)
    
    # URL tablosu
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'URL Pattern'
    hdr_cells[1].text = 'View'
    hdr_cells[2].text = 'Açıklama'
    
    url_data = [
        ['/', 'HomeView', 'Ana sayfa'],
        ['/products/', 'ProductListView', 'Ürün listesi'],
        ['/products/<slug>/', 'ProductDetailView', 'Ürün detay'],
        ['/cart/', 'CartView', 'Sepet görüntüle'],
        ['/cart/add/', 'AddToCartView', 'Sepete ekle'],
        ['/checkout/', 'CheckoutView', 'Ödeme sayfası'],
        ['/orders/', 'OrderListView', 'Sipariş geçmişi'],
        ['/auth/login/', 'LoginView', 'Giriş yap'],
        ['/auth/register/', 'RegisterView', 'Kayıt ol'],
        ['/admin/', 'admin.site.urls', 'Yönetim paneli'],
    ]
    
    for url_info in url_data:
        row_cells = table.add_row().cells
        row_cells[0].text = url_info[0]
        row_cells[1].text = url_info[1]
        row_cells[2].text = url_info[2]
    
    # TEST YAZIMI
    doc.add_heading('🧪 7. TEST YAZIMI VE KALİTE KONTROL', level=1)
    
    doc.add_heading('🎯 Ne Yapacaksın?', level=2)
    p2 = doc.add_paragraph('Kodunun doğru çalıştığından emin olmak için kapsamlı testler yaz.')
    
    doc.add_heading('📝 Test Türleri:', level=3)
    doc.add_paragraph('• Unit Tests (Model testleri)', style='List Bullet')
    doc.add_paragraph('• Integration Tests (View testleri)', style='List Bullet')
    doc.add_paragraph('• Functional Tests (Selenium ile)', style='List Bullet')
    doc.add_paragraph('• API Tests (DRF için)', style='List Bullet')
    doc.add_paragraph('• Performance Tests', style='List Bullet')
    
    doc.add_heading('💡 Örnek Test Kodu:', level=3)
    
    code_para = doc.add_paragraph()
    code_para.add_run('''
from django.test import TestCase, Client
from django.contrib.auth.models import User
from django.urls import reverse
from myapp.models import Product, Category

class ProductModelTest(TestCase):
    def setUp(self):
        self.category = Category.objects.create(
            name="Test Kategori",
            slug="test-kategori"
        )
        self.product = Product.objects.create(
            name="Test Ürün",
            slug="test-urun",
            description="Test açıklaması",
            price=99.99,
            stock_quantity=10,
            category=self.category
        )
    
    def test_product_creation(self):
        self.assertEqual(self.product.name, "Test Ürün")
        self.assertTrue(self.product.is_in_stock())
    
    def test_product_str_representation(self):
        self.assertEqual(str(self.product), "Test Ürün")

class ProductViewTest(TestCase):
    def setUp(self):
        self.client = Client()
        self.category = Category.objects.create(
            name="Test Kategori",
            slug="test-kategori"
        )
        self.product = Product.objects.create(
            name="Test Ürün",
            slug="test-urun",
            description="Test açıklaması",
            price=99.99,
            category=self.category
        )
    
    def test_product_list_view(self):
        response = self.client.get('/products/')
        self.assertEqual(response.status_code, 200)
        self.assertContains(response, "Test Ürün")
    
    def test_product_detail_view(self):
        url = reverse('product_detail', args=[self.product.slug])
        response = self.client.get(url)
        self.assertEqual(response.status_code, 200)
        self.assertContains(response, self.product.name)

# Test çalıştırma komutları:
# python manage.py test                    # Tüm testleri çalıştır
# python manage.py test myapp             # Belirli app testleri
# python manage.py test myapp.tests.ProductModelTest  # Belirli test sınıfı
# coverage run --source='.' manage.py test   # Coverage ile test
''').font.name = 'Courier New'
    
    # DEPLOYMENT
    doc.add_heading('🚀 8. DEPLOYMENT VE YAYIN', level=1)
    
    doc.add_heading('🎯 Ne Yapacaksın?', level=2)
    p3 = doc.add_paragraph('Projeyi canlı sunucuya yayınla ve kullanıcıların erişimine aç.')
    
    doc.add_heading('📝 Deployment Adımları:', level=3)
    
    doc.add_heading('🔹 1. Production Ayarları:', level=4)
    doc.add_paragraph('• DEBUG = False yap', style='List Bullet')
    doc.add_paragraph('• SECRET_KEY\'i çevre değişkeninde tut', style='List Bullet')
    doc.add_paragraph('• ALLOWED_HOSTS ayarla', style='List Bullet')
    doc.add_paragraph('• Static files yapılandırması', style='List Bullet')
    doc.add_paragraph('• Database ayarları (PostgreSQL)', style='List Bullet')
    doc.add_paragraph('• Email backend ayarları', style='List Bullet')
    
    doc.add_heading('🔹 2. Server Seçenekleri:', level=4)
    doc.add_paragraph('• Heroku (Kolay başlangıç)', style='List Bullet')
    doc.add_paragraph('• DigitalOcean (Güçlü ve ekonomik)', style='List Bullet')
    doc.add_paragraph('• AWS (Kurumsal çözüm)', style='List Bullet')
    doc.add_paragraph('• VPS (Kendi sunucun)', style='List Bullet')
    
    doc.add_heading('🔹 3. Heroku Deployment:', level=4)
    
    code_para2 = doc.add_paragraph()
    code_para2.add_run('''
# 1. Gerekli dosyaları oluştur:
# requirements.txt
pip freeze > requirements.txt

# Procfile
echo "web: gunicorn myproject.wsgi" > Procfile

# runtime.txt
echo "python-3.11.0" > runtime.txt

# 2. Git repository hazırla:
git init
git add .
git commit -m "Initial commit"

# 3. Heroku CLI ile deploy:
heroku create myproject-name
heroku addons:create heroku-postgresql:mini
heroku config:set SECRET_KEY="your-secret-key"
heroku config:set DEBUG=False
git push heroku main
heroku run python manage.py migrate
heroku run python manage.py collectstatic --noinput
heroku run python manage.py createsuperuser
''').font.name = 'Courier New'
    
    # İZLEME VE BAKIM
    doc.add_heading('📊 9. İZLEME VE BAKIM', level=1)
    
    doc.add_heading('🎯 Ne Yapacaksın?', level=2)
    p4 = doc.add_paragraph('Projeyi sürekli izle, güncellemeler yap ve sorunları çöz.')
    
    doc.add_heading('📝 İzleme Araçları:', level=3)
    doc.add_paragraph('• Error tracking (Sentry)', style='List Bullet')
    doc.add_paragraph('• Performance monitoring (New Relic)', style='List Bullet')
    doc.add_paragraph('• Uptime monitoring (UptimeRobot)', style='List Bullet')
    doc.add_paragraph('• Log analizi (Papertrail)', style='List Bullet')
    doc.add_paragraph('• User analytics (Google Analytics)', style='List Bullet')
    
    doc.add_heading('📝 Bakım Görevleri:', level=3)
    doc.add_paragraph('• Güvenlik güncellemeleri', style='List Bullet')
    doc.add_paragraph('• Database optimizasyonu', style='List Bullet')
    doc.add_paragraph('• Backup stratejisi', style='List Bullet')
    doc.add_paragraph('• Cache stratejisi', style='List Bullet')
    doc.add_paragraph('• CDN kullanımı', style='List Bullet')
    
    # SON CHECKLIST
    doc.add_heading('✅ 10. PROJE TAMAMLAMA CHECKLİSTİ', level=1)
    
    doc.add_heading('🔍 Kalite Kontrol:', level=3)
    checklist = [
        '☐ Tüm testler geçiyor',
        '☐ Code coverage %80 üzeri',
        '☐ Security scan temiz',
        '☐ Performance testleri yapıldı',
        '☐ Cross-browser uyumlu',
        '☐ Mobile responsive',
        '☐ SEO optimizasyonu',
        '☐ Accessibility (WCAG) uyumu',
        '☐ Error handling tamamlandı',
        '☐ Logging yapılandırıldı'
    ]
    
    for item in checklist:
        doc.add_paragraph(item, style='List Bullet')
    
    doc.add_heading('📚 Dokümantasyon:', level=3)
    doc_checklist = [
        '☐ README.md dosyası',
        '☐ API dokümantasyonu',
        '☐ Kurulum kılavuzu',
        '☐ Kullanıcı kılavuzu',
        '☐ Developer dokümantasyonu',
        '☐ Deployment kılavuzu',
        '☐ Troubleshooting rehberi',
        '☐ Changelog dosyası',
        '☐ License bilgisi',
        '☐ Contributing guidelines'
    ]
    
    for item in doc_checklist:
        doc.add_paragraph(item, style='List Bullet')
    
    doc.add_heading('🎉 TEBRİKLER!', level=1)
    congrats = doc.add_paragraph('Django projen başarıyla tamamlandı! Artık profesyonel bir web uygulamasına sahipsin. ')
    congrats.add_run('Unutma: Geliştirme süreci hiç bitmez. Kullanıcı geri bildirimlerini dinle, sürekli iyileştirmeler yap ve teknolojik gelişmeleri takip et.')
    
    doc.add_heading('🔄 Sonraki Adımlar:', level=3)
    doc.add_paragraph('• Kullanıcı geri bildirimlerini topla', style='List Bullet')
    doc.add_paragraph('• Analytics verilerini analiz et', style='List Bullet')
    doc.add_paragraph('• Yeni özellikler planla', style='List Bullet')
    doc.add_paragraph('• Teknik borçları temizle', style='List Bullet')
    doc.add_paragraph('• Kod refactoring yap', style='List Bullet')
    doc.add_paragraph('• Yeni teknolojileri araştır', style='List Bullet')
    
    # Dosyayı kaydet
    doc.save(r'c:\Users\ahmet.yildirir\Desktop\Apps\ProjectsDoc\Django_Kilavuz_Bolum4.docx')
    print("Django Kılavuzu Bölüm 4 oluşturuldu!")

if __name__ == "__main__":
    create_django_guide_part4()
